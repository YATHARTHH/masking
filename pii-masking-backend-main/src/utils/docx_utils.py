import os
import json
import cv2
import numpy as np
from PIL import Image
import fitz
import easyocr
import textdistance
from docx import Document

from google import genai
from google.genai import types

UPLOAD_FOLDER = "uploads"
PROCESSED_FOLDER = "processed"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
client = genai.Client(api_key=GEMINI_API_KEY)


def process_docx(docx_path,pii_category,highlight_mode):
    doc = Document(docx_path)
    print("Ye wala run kar raha hai")
    prompt = """Analyze the provided document and carefully identify any Personally Identifiable Information (PII) Which i am Mentioning. Your task is to detect, document, and categorize PII elements accurately without making any changes or masking.
        Strictly only Extract the Info Which comes under the PII Category i have Assigning
        Output Format:
            The output should be a detailed, structured summary of all detected PII in the document or image. Do not make any changes or alterations to the data. The format should include the type of PII, the exact text identified, and the context or description where it was found, ensuring that each category of PII is well-documented.

            Example:

            {
            "detected_pii": [
                {
                "type": "Full Name",
                "original_text": "John Doe"
                },
                {
                "type": "Social Security Number",
                "original_text": "123-45-6789"
                },
                {
                "type": "Personal Email Address",
                "original_text": "john.doe@example.com"
                }
            ]
            }
        Rules:
        Be as descriptive as possible when identifying and categorizing each piece of PII.
        Ensure that the context around each detected PII element is documented, describing where and how the PII is presented.
        ENsure that the PII of any language should listed.
        Only identify the presence of PII; do not alter, mask, or remove any details.
        Ensure the output is properly structured and JSON-compliant.

        Document Content:
    """
    docu=""
    for para in doc.paragraphs:
        docu+=para.text.strip()
    from src.utils.gemini_utils import generate_content_with_retry
    response = generate_content_with_retry(
            client,
            model="gemini-2.5-flash",
            contents=[prompt+docu+"\n\n\nPII Categories to Identify: "+pii_category, ]
        )
    if response and response.text:
        try:
            pii_data = json.loads(response.text[response.text.index("{"):response.text.rindex("}") + 1])
            from src.utils.gemini_utils import filter_pii_by_categories
            filtered_list = filter_pii_by_categories(pii_data.get("detected_pii", []), pii_category)
            print("Filtered PII:", filtered_list)
            l=[]
            for i in filtered_list:
                l.append((i["original_text"],i["type"]))
            for para in doc.paragraphs:

                if not para.text.strip():
                    continue
                for pii in l:
                    if pii[0] in para.text:
                        print(f"Detected PII in: {para.text}")
                        if highlight_mode=="named_replacement":
                            para.text = para.text.replace(pii[0], pii[1])
                        elif highlight_mode=="replacement":
                            para.text = para.text.replace(pii[0], "[MASKED]")
                        else:
                            para.text = para.text.replace(pii[0], len(pii)*"X")
            document_output_path = os.path.join(PROCESSED_FOLDER, os.path.basename(docx_path))
            doc.save(document_output_path)   
            return document_output_path  
        except Exception as e:
            print(e)